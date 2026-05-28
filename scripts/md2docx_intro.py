# -*- coding: utf-8 -*-
"""把 docs/产品介绍与报价说明.md 转成 Word(.docx)。
针对该文档已知结构做最小化解析：H1/H2/H3、表格、无序列表(含一级缩进)、引用、加粗、分隔线。
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

SRC = sys.argv[1] if len(sys.argv) > 1 else r"E:/my/桌面软件tauri/knowledge_base/docs/产品介绍与报价说明.md"
OUT = sys.argv[2] if len(sys.argv) > 2 else SRC.rsplit(".", 1)[0] + ".docx"

EAST_ASIA_FONT = "微软雅黑"
LATIN_FONT = "Segoe UI"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = LATIN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_inline(paragraph, text, base_size=None, base_bold=False):
    """解析 **加粗** 行内标记，逐段写 run。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, bold=base_bold)


def strip_bold(text):
    return text.replace("**", "")


def main():
    with open(SRC, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    # 正文默认字体
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 分隔线
        if stripped == "---":
            i += 1
            continue

        # 标题
        if stripped.startswith("# "):
            p = doc.add_heading(level=0)
            add_inline(p, strip_bold(stripped[2:]))
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(level=1)
            add_inline(p, strip_bold(stripped[3:]))
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(level=2)
            add_inline(p, strip_bold(stripped[4:]))
            i += 1
            continue

        # 表格：连续以 | 开头的行
        if stripped.startswith("|"):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            # 解析单元格
            rows = []
            for tl in tbl_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            # 去掉分隔行(---|---)
            rows = [r for r in rows if not all(set(c) <= set("-: ") and c for c in r)]
            if rows:
                ncol = len(rows[0])
                table = doc.add_table(rows=0, cols=ncol)
                table.style = "Light Grid Accent 1"
                for ri, r in enumerate(rows):
                    cells = table.add_row().cells
                    for ci in range(ncol):
                        txt = r[ci] if ci < len(r) else ""
                        cell = cells[ci]
                        cell.paragraphs[0].text = ""
                        add_inline(cell.paragraphs[0], txt, base_bold=(ri == 0))
            continue

        # 引用块
        if stripped.startswith(">"):
            quote = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run_mark = p.add_run("")
            add_inline(p, quote)
            for r in p.runs:
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
            i += 1
            continue

        # 列表项
        m = re.match(r"^(\s*)- (.*)$", line)
        if m:
            indent = len(m.group(1))
            content = m.group(2)
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print("OK ->", OUT)


if __name__ == "__main__":
    main()
